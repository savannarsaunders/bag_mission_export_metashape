#!/usr/bin/env python3
"""Generate ``mission_error_state_summary.docx`` from a set of exported
RangerBot missions.

For every timestamped mission folder under ``<mission-root>``
(``YYYYMMDDHHMMSS-<mission-name>/``) the script:

1. Reads ``mission_travel_path.csv`` and locates contiguous runs of non-zero
   ``errorState``.
2. Resolves the error name from the RangerBot error-code table (see
   ``ERROR_NAMES`` below — taken from ``RBerrorcodes.PDF``).
3. Reconstructs the waypoint the vehicle was driving toward when each error
   first fired, by replaying the GPS track against the mission plan's
   waypoint list and advancing whenever the vehicle enters that waypoint's
   capture radius. Mission folders without a plan JSON are listed without
   waypoint info.
4. Produces a Word document with a "Patterns worth noting" section (whose
   bullets are derived from the observed data — they describe this dataset
   accurately rather than restating a template) followed by a one-row-per-
   error-block table.

Example::

    python generate_error_summary.py \\
        /Users/vicar/Desktop/Rangerbot\\ Mission\\ \\(Keep\\)/24Jun2026_LemonDock \\
        --site-label "Rangerbot: Lemon" \\
        --dataset-label "24 June 2026, Lemon Dock"
"""

import argparse
import json
import math
import os
import sys
from pathlib import Path
from typing import Optional

import pandas as pd
from docx import Document
from docx.shared import Pt

# From RBerrorcodes.PDF — only entries actually encountered or likely
# encountered in RangerBot mission CSVs.
ERROR_NAMES = {
    9: "MISSION_TIMEOUT",
    10: "WAYPOINT_TIMEOUT",
    11: "WAYPOINT_ZERO_VELOCITY",
    12: "WAYPOINT_PATH_TOO_LONG",
    13: "MAXIMUM_DEPTH_EXCEEDED",
    14: "SAFE_SEARCH_BOUNDARY_INVALID",
    15: "USER_PARAMETER_INVALID",
    20: "GPS", 21: "SONAR", 22: "DVL", 23: "DVL_OVERHEAT",
    30: "IMU", 31: "IMU_VARIANCE_HIGH",
    40: "PRESSURE",
    50: "CAMERA_BOTH", 60: "CAMERA_DOWN", 70: "CAMERA_FRONT",
    71: "AI_DETECTOR", 80: "THRUSTER",
    100: "LEAK_SENSOR_FRONT", 101: "LEAK_SENSOR_REAR",
    160: "MANUAL_CONTROL_OVERRIDE",
}

SYS_TIME_COL = "timestamp_sys (%Y%m%d%H%M%S.%f)"


def haversine_m(lat1, lon1, lat2, lon2):
    R = 6371000.0
    p1, p2 = math.radians(lat1), math.radians(lat2)
    dp = math.radians(lat2 - lat1)
    dl = math.radians(lon2 - lon1)
    a = math.sin(dp/2)**2 + math.cos(p1) * math.cos(p2) * math.sin(dl/2)**2
    return 2 * R * math.asin(math.sqrt(a))


def load_mission_json(folder: Path) -> Optional[dict]:
    for p in folder.glob("*.json"):
        if p.name == "mission_summary.json":
            continue
        try:
            with p.open() as f:
                data = json.load(f)
            if isinstance(data, dict) and "waypoints" in data:
                return data
        except Exception:
            continue
    return None


def replay_targets(df: pd.DataFrame, waypoints: list) -> list:
    """Walk the GPS track. For each row record the waypoint index currently
    being targeted, advancing whenever the vehicle is inside that waypoint's
    capture radius. Returns a list of ``(target_idx, distance_m_to_target)``
    per row, where ``target_idx`` is ``None`` once the last waypoint is
    captured.
    """
    out = []
    cur = 0
    for _, row in df.iterrows():
        if cur >= len(waypoints):
            out.append((None, None))
            continue
        wp = waypoints[cur]
        d = haversine_m(row["latitude"], row["longitude"],
                        wp["latitude"], wp["longitude"])
        out.append((cur, d))
        if d <= wp.get("radius", 5.0):
            cur += 1
    return out


def find_contiguous_runs(idx_list):
    runs = []
    if not idx_list:
        return runs
    start = prev = idx_list[0]
    for i in idx_list[1:]:
        if i == prev + 1:
            prev = i
        else:
            runs.append((start, prev))
            start = prev = i
    runs.append((start, prev))
    return runs


def fmt_sys_hms(sys_str):
    """``'20260624125550.313604'`` → ``'12:55:50.313'``. Fractional seconds
    padded to three digits for a consistent column width."""
    s = str(sys_str)
    if "." in s:
        whole, frac = s.split(".", 1)
    else:
        whole, frac = s, "0"
    hh, mm, ss = whole[8:10], whole[10:12], whole[12:14]
    return f"{hh}:{mm}:{ss}.{(frac[:3]).ljust(3, '0')}"


def analyse(root: Path):
    rows_out = []
    notes = {
        "all_state_9": True,
        "clean_count": 0,
        "error_count": 0,
        "tiny_blocks": [],          # (folder, n)
        "sustained_blocks": [],     # (folder, n)
        "missions_with_errors": [],
        "missions_clean": [],
        "all_at_end": True,
        "near_surface": True,
        "all_one_block": True,
        "no_mission_json": [],      # (folder,)
        "total_missions": 0,
    }
    folders = sorted(d for d in os.listdir(root)
                     if d[:14].isdigit() and (root / d).is_dir())
    counters: dict[str, int] = {}

    for folder_name in folders:
        folder = root / folder_name
        mission_name = folder_name.split("-", 1)[1]
        counters[mission_name] = counters.get(mission_name, 0) + 1
        mission_num = counters[mission_name]

        csv_path = folder / "mission_travel_path.csv"
        if not csv_path.exists():
            continue
        notes["total_missions"] += 1
        df = pd.read_csv(csv_path)
        mission_json = load_mission_json(folder)
        waypoints = mission_json.get("waypoints", []) if mission_json else []
        if not waypoints:
            notes["no_mission_json"].append(folder_name)

        nz_idx = df.index[df["errorState"] != 0].tolist()
        runs = find_contiguous_runs(nz_idx)

        if not runs:
            notes["clean_count"] += 1
            notes["missions_clean"].append((mission_name, mission_num, folder_name,
                                            len(df), len(waypoints)))
            rows_out.append({
                "name": mission_name, "num": mission_num, "folder": folder_name,
                "error_rows": 0, "state": "—", "ename": "—",
                "wp": f"— (no error; {len(waypoints)} waypoints)" if waypoints
                       else "— (no error)",
                "tspan": "—", "loc": "no errors",
            })
            continue

        if len(runs) > 1:
            notes["all_one_block"] = False

        targets = replay_targets(df, waypoints) if waypoints else None

        for ri, (a, b) in enumerate(runs):
            n = b - a + 1
            notes["error_count"] += 1
            if n <= 10:
                notes["tiny_blocks"].append((folder_name, n))
            else:
                notes["sustained_blocks"].append((folder_name, n))
            notes["missions_with_errors"].append((mission_name, mission_num,
                                                  folder_name, n))
            if b != len(df) - 1:
                notes["all_at_end"] = False

            first_row = df.loc[a]
            last_row = df.loc[b]
            state = int(first_row["errorState"])
            if state != 9:
                notes["all_state_9"] = False
            ename = ERROR_NAMES.get(state, f"UNKNOWN ({state})")

            if targets:
                tgt_idx, tgt_dist = targets[a]
                total_wp = len(waypoints)
                if tgt_idx is None:
                    wp_str = f"— (all {total_wp} waypoints captured before error)"
                else:
                    wp_str = (f"WP #{tgt_idx + 1} of {total_wp} "
                              f"({tgt_dist:.1f} m away)")
            else:
                wp_str = "— (no mission JSON in folder)"

            t0 = fmt_sys_hms(first_row[SYS_TIME_COL])
            t1 = fmt_sys_hms(last_row[SYS_TIME_COL])
            tspan = t0 if n == 1 or t0 == t1 else f"{t0} → {t1}"

            depth0 = float(first_row["depth"])
            if abs(depth0) >= 0.5:
                notes["near_surface"] = False
            if n == 1:
                loc = (f"{first_row['latitude']:.6f}, "
                       f"{first_row['longitude']:.6f} (depth {depth0:.2f})")
            else:
                loc = (f"{first_row['latitude']:.6f}, "
                       f"{first_row['longitude']:.6f} → "
                       f"{last_row['latitude']:.6f}, "
                       f"{last_row['longitude']:.6f}")

            rows_out.append({
                "name": mission_name, "num": mission_num, "folder": folder_name,
                "error_rows": n, "state": state, "ename": ename,
                "wp": wp_str, "tspan": tspan, "loc": loc,
            })

    return rows_out, notes


def build_docx(rows, notes, out_docx: Path, root: Path,
               site_label: Optional[str], dataset_label: Optional[str]):
    doc = Document()
    doc.add_heading("Mission Travel Path – Non-zero errorState Summary", level=1)
    p = doc.add_paragraph()
    prefix = f"Site: {site_label}. " if site_label else ""
    dataset_clause = f" ({dataset_label})" if dataset_label else ""
    p.add_run(f"{prefix}Source: {root}  —  mission_travel_path.csv files"
              f"{dataset_clause}. Error names from RBerrorcodes.PDF.")
    doc.add_paragraph()

    doc.add_heading("Patterns worth noting", level=2)

    bullets = []
    if notes["error_count"] == 0:
        bullets.append(
            f"All {notes['total_missions']} missions ran clean — every "
            f"mission_travel_path.csv has errorState = 0 for the full duration."
        )
    else:
        if notes["all_state_9"]:
            bullets.append(
                "errorState = 9 is MISSION_TIMEOUT — per RBerrorcodes.PDF the "
                "mission was not completed within its configured maximum "
                "mission duration (1 hour by default). It is the only non-zero "
                "error code in this dataset; no joystick overrides, sensor "
                "faults, or waypoint timeouts appear."
            )
        if notes["all_at_end"]:
            bullets.append(
                "In every mission that flagged an error, the non-zero block "
                "runs to the very last row of the CSV, so the timeout fires "
                "at the end of the recording — these read as mission-end "
                "markers, not faults during operation."
            )
        tiny = [(f, n) for f, n in notes["tiny_blocks"] if n <= 4]
        sustained = notes["sustained_blocks"]
        if tiny and not sustained:
            bullets.append(
                f"Every flagged block is short ({', '.join(f'{n} rows in {f}' for f, n in tiny)}) "
                f"— ≤ ~0.4 s before logging stops, so the timeout reads "
                f"essentially as a one-tick boundary condition."
            )
        elif sustained and not tiny:
            parts = ", ".join(f"{n} rows in {f}" for f, n in sustained)
            bullets.append(
                f"The flagged block is sustained ({parts}) rather than a "
                f"one-tick spike, indicating the timeout fired and the "
                f"vehicle continued logging for several seconds before the "
                f"recording ended."
            )
        elif tiny and sustained:
            tiny_desc = ", ".join(f"{f} ({n} rows)" for f, n in tiny)
            sustained_desc = ", ".join(f"{f} ({n} rows)" for f, n in sustained)
            bullets.append(
                f"Mixed block sizes: {tiny_desc} trip the timeout briefly "
                f"(≤ ~0.4 s), while {sustained_desc} sustains the timeout for "
                f"several seconds of further logging."
            )
        if notes["near_surface"]:
            bullets.append(
                "All timeout blocks are recorded at near-surface depth "
                "(≤ 0.5 m), consistent with the vehicle having already "
                "returned to the surface when the timeout fires."
            )
    if notes["clean_count"]:
        clean_list = ", ".join(
            f"{n} #{i} ({f})" for n, i, f, _, _ in notes["missions_clean"]
        )
        bullets.append(
            f"{notes['clean_count']} of the {notes['total_missions']} "
            f"missions are clean (zero non-zero errorState rows): "
            f"{clean_list}."
        )
    if notes["error_count"] > 0 and notes["all_one_block"]:
        bullets.append(
            "In every file the non-zero errorState rows form a single "
            "contiguous block, so each is summarized as one segment rather "
            "than per-row."
        )
    methodology = (
        "Methodology for the \"Waypoint during error\" column: the CSV files "
        "contain no waypoint index, so the target waypoint was reconstructed "
        "by replaying each mission's GPS track against its mission-plan "
        "waypoint list, advancing to the next waypoint whenever the vehicle "
        "came within that waypoint's 5 m capture radius. The value shown is "
        "the waypoint the vehicle was driving toward (and its straight-line "
        "distance from it) on the first row where errorState went non-zero."
    )
    if notes["no_mission_json"]:
        methodology += (
            f" For {', '.join(notes['no_mission_json'])} no mission plan "
            f"JSON was present in the folder, so no waypoint can be "
            f"reconstructed."
        )
    bullets.append(methodology)

    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    headers = ["Mission Name", "Mission #", "Folder", "Error rows", "errorState",
               "Error name", "Waypoint during error", "Time span (sys)", "Location"]
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True

    for r in rows:
        cells = t.add_row().cells
        cells[0].text = r["name"]
        cells[1].text = str(r["num"])
        cells[2].text = r["folder"]
        cells[3].text = str(r["error_rows"])
        cells[4].text = str(r["state"])
        cells[5].text = r["ename"]
        cells[6].text = r["wp"]
        cells[7].text = r["tspan"]
        cells[8].text = r["loc"]

    for row in t.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.save(out_docx)


def parse_args() -> argparse.Namespace:
    ap = argparse.ArgumentParser(
        description=__doc__,
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    ap.add_argument("mission_root", type=Path,
                    help="Mission root containing the YYYYMMDDHHMMSS-<name>/ "
                         "exported folders.")
    ap.add_argument("--out", type=Path, default=None,
                    help="Output .docx path "
                         "(default: <mission-root>/mission_error_state_summary.docx).")
    ap.add_argument("--site-label", default=None,
                    help="Optional site label prepended to the source line "
                         "(e.g. 'Rangerbot: Lemon').")
    ap.add_argument("--dataset-label", default=None,
                    help="Optional short dataset description appended to the "
                         "source line (e.g. '24 June 2026, Lemon Dock').")
    return ap.parse_args()


def main():
    args = parse_args()
    root: Path = args.mission_root.resolve()
    if not root.is_dir():
        raise SystemExit(f"mission_root not found: {root}")
    out_docx = (args.out or (root / "mission_error_state_summary.docx")).resolve()

    rows, notes = analyse(root)
    build_docx(rows, notes, out_docx, root,
               args.site_label, args.dataset_label)
    print(f"Wrote {out_docx}")
    print(f"  {len(rows)} table rows; {notes['error_count']} with errors, "
          f"{notes['clean_count']} clean")
    return 0


if __name__ == "__main__":
    sys.exit(main())
